import os
import tkinter as tk
from tkinter import filedialog
from docx import Document


class Refchecker(tk.Tk):
    """Thanks to David Love and his book "Python Tkinter By Example"
    """

    def __init__(self):
        super().__init__()

        self.title('Check Document References')
        self.geometry('1000x600')

        self.create_file_frame()

        self.create_folder_frame()

        self.create_result_canvas()

        # self.create_header_frame()

    def create_file_frame(self):
        self.file_frame = tk.Frame(self)
        self.file_frame.place(
            relx=0.5, rely=0.1, relwidth=0.75, relheight=0.1, anchor='n')

        self.doc_entry = tk.Entry(self.file_frame, font=40)
        self.doc_entry.place(relwidth=0.65, relheight=.5)
        self.doc_entry.focus_set()

        self.browse_file_button = tk.Button(
            self.file_frame, text="Browse", font=40,
            command=self.get_file_name)
        self.browse_file_button.place(relx=0.7, relheight=.5, relwidth=0.1)

        self.check_file_button = tk.Button(
            self.file_frame, text="Check Document", font=40,
            command=self.run_document_check)
        self.check_file_button.place(relx=0.8, relheight=.5, relwidth=0.2)

    def create_folder_frame(self):
        self.folder_frame = tk.Frame(self)
        self.folder_frame.place(
            relx=0.5, rely=0.2, relwidth=0.75, relheight=0.1, anchor='n')

        self.folder_entry = tk.Entry(self.folder_frame, font=20)
        self.folder_entry.place(relwidth=0.65, relheight=.5)
        self.folder_entry.focus_set()

        self.browse_folder_button = tk.Button(
            self.folder_frame, text="Browse", font=20,
            command=self.get_folder_name)
        self.browse_folder_button.place(relx=0.7, relheight=.5, relwidth=0.1)

        self.check_file_button = tk.Button(
            self.folder_frame, text="Check Folder", font=20,
            command=self.run_folder_check)
        self.check_file_button.place(relx=0.8, relheight=.5, relwidth=0.2)

    # This was just if I liked having the header "frozen" vs. scrolling
    # def create_header_frame(self):
    #     self.header_frame = tk.Frame(self, bd=10)
    #     self.header_frame.place(relx=0.5, rely=0.2, relwidth=0.75,
    #                     relheight=0.1, anchor='n')

    #     self.doc_header = tk.Label(
    #         self.header_frame, text='Document Searched', bd=20)
    #     self.doc_header.grid(row=0, column=0)
    #     self.ref_header = tk.Label(
    #         self.header_frame, text='Document Referenced')
    #     self.ref_header.grid(row=0, column=1, ipadx=30)
    #     self.used_header = tk.Label(
    #         self.header_frame, text='Reference Used in Document')
    #     self.used_header.grid(row=0, column=2, ipadx=30)

    def create_result_canvas(self):
        self.result_canvas = tk.Canvas(self)

        self.result_frame = tk.Frame(self.result_canvas, bd=10)

        self.scrollbar = tk.Scrollbar(
            self.result_canvas, orient='vertical', command=self.result_canvas.yview)

        self.result_canvas.configure(yscrollcommand=self.scrollbar.set)

        self.result_canvas.place(
            relx=0.5, rely=0.3, relwidth=0.75, relheight=0.9, anchor='n')
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.canvas_frame = self.result_canvas.create_window((300, 0), window=self.result_frame, anchor='n')

        self.doc_header = tk.Label(
            self.result_frame, text='Document Searched', bd=20)
        self.doc_header.grid(row=0, column=0)
        self.ref_header = tk.Label(
            self.result_frame, text='Document Referenced')
        self.ref_header.grid(row=0, column=1, ipadx=30)
        self.used_header = tk.Label(
            self.result_frame, text='Reference Used in Document')
        self.used_header.grid(row=0, column=2, ipadx=30)

        self.bind("<Return>", self.run_document_check)
        self.bind("<Configure>", self.on_frame_configure)

    def get_file_name(self):
        self.file_name = tk.filedialog.askopenfilename(initialdir="/", title="Select file", filetypes=(("docx files", "*.docx"), ("doc files", "*.doc"), ("all files", "*.*")))
        self.doc_entry.insert(0, self.file_name)

    def get_folder_name(self):
        self.folder_name = tk.filedialog.askdirectory()
        self.folder_entry.insert(0, self.folder_name)

    def run_document_check(self):
        document = Document(self.doc_entry.get())

        doc_dict = {self.doc_entry.get(): self.check_document_references(document)}

        # references = self.check_document_references(document)

        self.display_result_grid(doc_dict)

    def run_folder_check(self):
        doc_list = os.listdir(self.folder_entry.get())

        doc_dict = {}

        for doc in doc_list:
            if os.path.splitext(doc)[-1] == '.doc' or os.path.splitext(doc)[-1] == '.docx':
                doc_dict[doc] = self.check_document_references(Document(os.path.join(self.folder_entry.get(), doc)))

        self.display_result_grid(doc_dict)

    def check_document_references(self, document):

        # TODO: Use the is_active via web service call

        references = {}

        for t in reversed(document.tables):
            if (t.columns[0].cells[0].text.upper().strip() == 'ID' or t.columns[0].cells[0].text.upper().strip() == 'DOCUMENT NUMBER') and (t.columns[1].cells[0].text.upper().strip() == 'TITLE' or t.columns[1].cells[0].text.upper().strip() == 'DOCUMENT TITLE'):
                references = {c.text.upper().strip(): {'is_referenced': 'No', 'is_active': 'No'}
                              for c in t.columns[0].cells[1:] if c.text.strip()}
                break

        if references:

            for ref in references:
                for p in document.paragraphs:
                    if ref in p.text:
                        references[ref]['is_referenced'] = 'Yes'
                        break

                for t in document.tables:
                    if (t.columns[0].cells[0].text.upper().strip() == 'ID' or t.columns[0].cells[0].text.upper().strip() == 'DOCUMENT NUMBER') and (t.columns[1].cells[0].text.upper().strip() == 'TITLE' or t.columns[1].cells[0].text.upper().strip() == 'DOCUMENT TITLE'):
                        continue

                    for r in t.rows:
                        for c in r.cells:
                            if ref in c.text.upper().strip():
                                references[ref]['is_referenced'] = 'Yes'

        else:
            references = {'No References Found': {
                'is_referenced': 'N/A', 'is_active': 'N/A'}}

        return references

    def display_result_grid(self, doc_dict):
        self.result_frame.grid_remove()
        row_count = 0
        for doc_key in doc_dict:
            for ref_key in doc_dict[doc_key]:
                self.doc_lbl = tk.Label(
                    self.result_frame, text=os.path.basename(doc_key))
                self.doc_lbl.grid(row=row_count+1, column=0)
                self.ref_lbl = tk.Label(self.result_frame, text=ref_key)
                self.ref_lbl.grid(row=row_count+1, column=1, ipadx=30)
                self.used_lbl = tk.Label(
                    self.result_frame, text=doc_dict[doc_key][ref_key]['is_referenced'])
                self.used_lbl.grid(row=row_count+1, column=2, ipadx=30)
                row_count += 1

    def on_frame_configure(self, event):
        self.result_canvas.configure(scrollregion=self.result_canvas.bbox("all"))


if __name__ == '__main__':
    refchecker = Refchecker()
    refchecker.mainloop()
